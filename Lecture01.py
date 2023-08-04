import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
import math
import os
from docx import Document
from docx.shared import Inches

def read_data_from_list():
    data = entry_data.get()
    try:
        data_list = list(map(float, data.split(',')))
        calculate_statistics(data_list)
    except ValueError:
        result_text.set("Invalid input. Please enter a valid list of numbers separated by commas.")

def read_data_from_excel():
    script_dir = os.path.dirname(__file__)
    file_path = os.path.join(script_dir, "lecture01_sample_data.xlsx")
    description_text = ("Please make sure the Excel file is formatted as follows:\n\n"
                        "1. The data should be in the first column.\n"
                        "2. There should be no headers in the first row.\n"
                        "3. The file should be in the same directory as this program.\n\n"
                        "Once you have the correct Excel file, click 'Open' to continue.")
    
    # Show the description pop-up window
    description_window = tk.Toplevel(root)
    description_window.title("Excel File Format")
    description_label = tk.Label(description_window, text=description_text, wraplength=400, justify="left")
    description_label.pack(padx=10, pady=10)
    
    # Open the Excel file
    file_path = filedialog.askopenfilename()
    description_window.destroy()  # Close the description window
    if file_path:
        try:
            data = pd.read_excel(file_path)
            data_list = data.iloc[:, 0].dropna().tolist()
            entry_data.delete(0, tk.END)
            entry_data.insert(tk.END, ",".join(str(item) for item in data_list))
            calculate_statistics(data_list)
        except Exception:
            result_text.set("Error reading data from the selected Excel file.")

def open_excel_file():
    script_dir = os.path.dirname(__file__)
    file_path = os.path.join(script_dir, "lecture01_sample_data.xlsx")
    try:
        os.startfile(file_path)
    except Exception:
        messagebox.showerror("Error", "Unable to open the Excel file.")

def calculate_statistics(data_list):
    n = len(data_list)
    mean = sum(data_list) / n
    data_list.sort()
    
    if n % 2 == 1:
        median = data_list[(n + 1) // 2 - 1]
    else:
        median = (data_list[n // 2 - 1] + data_list[n // 2]) / 2

    mode = max(data_list, key=data_list.count)

    data_range = max(data_list) - min(data_list)

    variance = sum((xi - mean) ** 2 for xi in data_list) / n

    standard_deviation = math.sqrt(variance)

    mad = sum(abs(xi - mean) for xi in data_list) / n

    q1 = data_list[int(n * 0.25)]
    q3 = data_list[int(n * 0.75)]
    iqr = q3 - q1

    coefficient_of_variation = (standard_deviation / mean) * 100

    skewness = sum((xi - mean) ** 3 for xi in data_list) / (n * standard_deviation ** 3)

    result_text.set(f"Mean (μ): {mean}\nMedian: {median}\nMode: {mode}\nRange: {data_range}"
                    f"\nVariance (σ²): {variance}\nStandard Deviation (σ): {standard_deviation}"
                    f"\nMean Absolute Deviation (MAD): {mad}\nInterquartile Range (IQR): {iqr}"
                    f"\nCoefficient of Variation (CV): {coefficient_of_variation:.2f}%"
                    f"\nSkewness: {skewness}")

    # Ask the user if they want to enter other values or save to DOCX file
    response = messagebox.askquestion("Save to DOCX?", "Do you want to save the results to a DOCX file?", icon='question')
    if response == 'yes':
        save_to_docx(data_list, mean, median, mode, data_range, variance, standard_deviation, mad, iqr, coefficient_of_variation, skewness)
    else:
        # Show the result on the screen and ask the user if they want to continue or quit
        response = messagebox.askquestion("Continue?", "Do you want to enter other values?", icon='question')
        if response == 'yes':
            entry_data.delete(0, tk.END)
            result_text.set("")
        else:
            root.destroy()

def save_to_docx(data_list, mean, median, mode, data_range, variance, standard_deviation, mad, iqr, coefficient_of_variation, skewness):
    # Create a new Word document
    doc = Document()

    # Add content to the document
    doc.add_heading('Statistics Results', level=1)

    # Add parsed data from the list or Excel sheet to the input section
    doc.add_heading('Input', level=2)
    doc.add_paragraph("Data: " + ", ".join(str(item) for item in data_list))

    # Add formulas used in calculations to the formula section
    doc.add_heading('Formulas Used', level=2)
    doc.add_paragraph("1. Mean (μ): μ = (Σ(xi)) / n")
    doc.add_paragraph("2. Median:")
    doc.add_paragraph("   - If n is odd: Median = Value at position (n + 1) / 2 in the ordered dataset.")
    doc.add_paragraph("   - If n is even: Median = Average of the values at positions n / 2 and (n / 2) + 1 in the ordered dataset.")
    doc.add_paragraph("3. Mode:")
    doc.add_paragraph("   - Mode is the value that appears most frequently in the dataset.")
    doc.add_paragraph("4. Range: Range = Maximum value - Minimum value")
    doc.add_paragraph("5. Variance (σ²): σ² = Σ((xi - μ)²) / n")
    doc.add_paragraph("6. Standard Deviation (σ): σ = √σ²")
    doc.add_paragraph("7. Mean Absolute Deviation (MAD): MAD = Σ(|xi - μ|) / n")
    doc.add_paragraph("8. Interquartile Range (IQR): IQR = Q3 - Q1")
    doc.add_paragraph("   - where Q1 is the first quartile (25th percentile) and Q3 is the third quartile (75th percentile).")
    doc.add_paragraph("9. Coefficient of Variation (CV): CV = (σ / μ) * 100")
    doc.add_paragraph("   - where σ is the standard deviation and μ is the mean.")
    doc.add_paragraph("10. Skewness: Skewness = (Σ((xi - μ)³) / (n * σ³))")
    doc.add_paragraph("    - where xi represents each data point, μ is the mean, σ is the standard deviation, and n is the total number of data points.")
    doc.add_paragraph("11. Outliers: Z-Scores: Z-Score = (Data point - Mean) / Standard Deviation")

    # Add the calculated statistics to the result section as a table
    doc.add_heading('Result', level=2)
    table = doc.add_table(rows=2, cols=7)
    table.style = 'Table Grid'
    table.autofit = True

    table.cell(0, 0).text = "Mean (μ)"
    table.cell(0, 1).text = "Median"
    table.cell(0, 2).text = "Mode"
    table.cell(0, 3).text = "Range"
    table.cell(0, 4).text = "Variance (σ²)"
    table.cell(0, 5).text = "Standard Deviation (σ)"
    table.cell(0, 6).text = "Mean Absolute Deviation (MAD)"

    table.cell(1, 0).text = str(mean)
    table.cell(1, 1).text = str(median)
    table.cell(1, 2).text = str(mode)
    table.cell(1, 3).text = str(data_range)
    table.cell(1, 4).text = str(variance)
    table.cell(1, 5).text = str(standard_deviation)
    table.cell(1, 6).text = str(mad)
    
    # Ask the user to provide the name and path to save the DOCX file
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])

    if file_path:
        doc.save(file_path)
        messagebox.showinfo("Success", "Results saved to the DOCX file successfully.")
    else:
        messagebox.showinfo("Info", "Results not saved to the DOCX file.")

    # Ask the user if they want to enter other values or quit
    response = messagebox.askquestion("Continue?", "Do you want to enter other values?", icon='question')
    if response == 'yes':
        entry_data.delete(0, tk.END)
        result_text.set("")
    else:
        root.destroy()
        
def show_about_page():
    about_window = tk.Toplevel(root)
    about_window.title("About Statistics Calculator")

    about_text = ("Statistics Calculator GUI\n\n"
                  "This program calculates various statistical measures for a given dataset. "
                  "You can enter data manually in the entry box using comma-separated values, "
                  "or load data from an Excel file. The program will display the following statistics:\n\n"
                  "1. Mean (μ): μ = (Σ(xi)) / n\n"
                  "2. Median:\n"
                  "   - If n is odd: Median = Value at position (n + 1) / 2 in the ordered dataset.\n"
                  "   - If n is even: Median = Average of the values at positions n / 2 and (n / 2) + 1 in the ordered dataset.\n"
                  "3. Mode:\n"
                  "   - Mode is the value that appears most frequently in the dataset.\n"
                  "4. Range: Range = Maximum value - Minimum value\n"
                  "5. Variance (σ²): σ² = Σ((xi - μ)²) / n\n"
                  "6. Standard Deviation (σ): σ = √σ²\n"
                  "7. Mean Absolute Deviation (MAD): MAD = Σ(|xi - μ|) / n\n"
                  "8. Interquartile Range (IQR): IQR = Q3 - Q1\n"
                  "   - where Q1 is the first quartile (25th percentile) and Q3 is the third quartile (75th percentile).\n"
                  "9. Coefficient of Variation (CV): CV = (σ / μ) * 100\n"
                  "   - where σ is the standard deviation and μ is the mean.\n"
                  "10. Skewness: Skewness = (Σ((xi - μ)³) / (n * σ³))\n"
                  "    - where xi represents each data point, μ is the mean, σ is the standard deviation, and n is the total number of data points.\n"
                  "11. Outliers: Z-Scores: Z-Score = (Data point - Mean) / Standard Deviation\n\n"
                  "Developed by Mahmoud Tawfeek\n"
                  "Version: 1.0\n"
                  "Date: 4 August 2023\n")

    about_label = tk.Label(about_window, text=about_text, wraplength=400, justify="left")
    about_label.pack(padx=10, pady=10)

def show_formulas_page():
    formulas_window = tk.Toplevel(root)
    formulas_window.title("Formulas Used")

    formulas_text = ("Formulas Used in Calculations\n\n"
                     "1. Mean (μ): μ = (Σ(xi)) / n\n"
                     "2. Median:\n"
                     "   - If n is odd: Median = Value at position (n + 1) / 2 in the ordered dataset.\n"
                     "   - If n is even: Median = Average of the values at positions n / 2 and (n / 2) + 1 in the ordered dataset.\n"
                     "3. Mode:\n"
                     "   - Mode is the value that appears most frequently in the dataset.\n"
                     "4. Range: Range = Maximum value - Minimum value\n"
                     "5. Variance (σ²): σ² = Σ((xi - μ)²) / n\n"
                     "6. Standard Deviation (σ): σ = √σ²\n"
                     "7. Mean Absolute Deviation (MAD): MAD = Σ(|xi - μ|) / n\n"
                     "8. Interquartile Range (IQR): IQR = Q3 - Q1\n"
                     "   - where Q1 is the first quartile (25th percentile) and Q3 is the third quartile (75th percentile).\n"
                     "9. Coefficient of Variation (CV): CV = (σ / μ) * 100\n"
                     "   - where σ is the standard deviation and μ is the mean.\n"
                     "10. Skewness: Skewness = (Σ((xi - μ)³) / (n * σ³))\n"
                     "    - where xi represents each data point, μ is the mean, σ is the standard deviation, and n is the total number of data points.\n"
                     "11. Outliers: Z-Scores: Z-Score = (Data point - Mean) / Standard Deviation\n\n")

    formulas_label = tk.Label(formulas_window, text=formulas_text, wraplength=400, justify="left")
    formulas_label.pack(padx=10, pady=10)

def load_sample_list():
    sample_list = [15, 23, 12, 34, 45, 67, 56, 78, 23, 10]
    entry_data.delete(0, tk.END)
    entry_data.insert(tk.END, ",".join(str(item) for item in sample_list))

def load_sample_excel():
    script_dir = os.path.dirname(__file__)
    file_path = os.path.join(script_dir, "lecture01_sample_data.xlsx")
    try:
        data = pd.read_excel(file_path)
        data_list = data.iloc[:, 0].dropna().tolist()
        entry_data.delete(0, tk.END)
        entry_data.insert(tk.END, ",".join(str(item) for item in data_list))
    except Exception:
        result_text.set("Error loading the sample Excel file.")

# Create the main GUI window
root = tk.Tk()
root.title("Statistics Calculator")

# Create widgets
label_data = tk.Label(root, text="Enter data (comma-separated):")
entry_data = tk.Entry(root, width=50)
button_about = tk.Button(root, text="About", command=show_about_page)
button_formulas = tk.Button(root, text="Formulas Used", command=show_formulas_page)
button_load_sample_list = tk.Button(root, text="Load Sample List", command=load_sample_list)
button_open_sample_excel = tk.Button(root, text="Open Sample Excel", command=open_excel_file)
button_load_sample_excel = tk.Button(root, text="Load Sample Excel", command=load_sample_excel)
button_excel = tk.Button(root, text="Calculate from Excel", command=read_data_from_excel)
button_list = tk.Button(root, text="Calculate from List", command=read_data_from_list)
result_text = tk.StringVar()
label_result = tk.Label(root, textvariable=result_text, wraplength=400, justify="left")

# Grid layout
label_data.grid(row=0, column=0, columnspan=2, padx=5, pady=5)
entry_data.grid(row=0, column=2, columnspan=2, padx=5, pady=5)
button_about.grid(row=1, column=0, padx=5, pady=5)
button_formulas.grid(row=1, column=1, padx=5, pady=5)
button_load_sample_list.grid(row=1, column=2, padx=5, pady=5)
button_open_sample_excel.grid(row=1, column=3, padx=5, pady=5)
button_load_sample_excel.grid(row=1, column=4, padx=5, pady=5)
button_excel.grid(row=2, column=2, padx=5, pady=5)
button_list.grid(row=2, column=3, padx=5, pady=5)
label_result.grid(row=4, column=0, columnspan=5, padx=5, pady=5)

root.mainloop()